#!/usr/bin/env python3
"""Generate AEO Audit Report for bellanico.com — Fresh Design Studio"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import sys
from pathlib import Path

# ── Colours ──────────────────────────────────────────────────────────────────
GREEN       = RGBColor(0x4C, 0x7C, 0x09)   # Fresh Design Studio brand green
DARK_GREEN  = RGBColor(0x2E, 0x4D, 0x04)
LIGHT_GREEN = RGBColor(0xF1, 0xF6, 0xE6)
DARK_GREY   = RGBColor(0x33, 0x33, 0x33)
MID_GREY    = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
RED         = RGBColor(0xC0, 0x39, 0x2B)
AMBER       = RGBColor(0xE6, 0x7E, 0x22)
BLUE        = RGBColor(0x1A, 0x5C, 0x8A)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='CCCCCC', size='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(doc, text='', style='Normal', bold=False, size=None,
             color=None, align=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size:  run.font.size  = Pt(size)
        if color: run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1, color=GREEN, size=None, space_before=18, space_after=6):
    sizes = {1: 20, 2: 15, 3: 13}
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.color.rgb = color
        if size: run.font.size = Pt(size)
        else:    run.font.size = Pt(sizes.get(level, 13))
    return p

def add_divider(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:color'), '4C7C09')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def make_label_cell(cell, text, bg: RGBColor, fg: RGBColor = WHITE):
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(8.5)
    run.font.color.rgb = fg

def add_finding_table(doc, findings):
    """findings = list of (area, issue, impact, fix_type, location_label)"""
    headers = ['Area', 'Finding', 'Impact on AI Visibility', 'Fix Type', 'Where']
    col_widths = [Cm(2.8), Cm(5.5), Cm(5.2), Cm(2.2), Cm(2.2)]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        set_cell_bg(cell, GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for idx, (area, issue, impact, fix_type, location) in enumerate(findings):
        row  = tbl.add_row()
        bg   = LIGHT_GREY if idx % 2 == 0 else WHITE

        # Area
        c = row.cells[0]; c.width = col_widths[0]
        set_cell_bg(c, LIGHT_GREEN)
        p = c.paragraphs[0]
        run = p.add_run(area); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREEN

        # Issue
        c = row.cells[1]; c.width = col_widths[1]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        run = p.add_run(issue); run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREY

        # Impact
        c = row.cells[2]; c.width = col_widths[2]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        run = p.add_run(impact); run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY

        # Fix Type badge
        c = row.cells[3]; c.width = col_widths[3]
        ft_lower = fix_type.lower()
        if 'critical' in ft_lower:
            make_label_cell(c, fix_type, RED)
        elif 'important' in ft_lower:
            make_label_cell(c, fix_type, AMBER, DARK_GREY)
        else:
            make_label_cell(c, fix_type, MID_GREY)

        # Location badge
        c = row.cells[4]; c.width = col_widths[4]
        if 'front' in location.lower():
            make_label_cell(c, location, BLUE)
        else:
            make_label_cell(c, location, DARK_GREY)

    doc.add_paragraph()
    return tbl

def add_rec_table(doc, recs):
    """recs = list of (no, recommendation, details, location)"""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(0.8), Cm(3.8), Cm(9.2), Cm(2.0)]

    headers = ['#', 'Recommendation', 'What To Do', 'Where']
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = hdr.cells[i]
        cell.width = w
        set_cell_bg(cell, DARK_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(9); run.font.color.rgb = WHITE

    for idx, (no, rec, detail, loc) in enumerate(recs):
        row = tbl.add_row()
        bg  = LIGHT_GREY if idx % 2 == 0 else WHITE

        c = row.cells[0]; c.width = widths[0]
        set_cell_bg(c, LIGHT_GREEN)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(no)); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREEN

        c = row.cells[1]; c.width = widths[1]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        run = p.add_run(rec); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREY

        c = row.cells[2]; c.width = widths[2]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        run = p.add_run(detail); run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY

        c = row.cells[3]; c.width = widths[3]
        if 'front' in loc.lower():
            make_label_cell(c, loc, BLUE)
        else:
            make_label_cell(c, loc, DARK_GREY)

    doc.add_paragraph()
    return tbl


# ── MAIN ──────────────────────────────────────────────────────────────────────
def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── COVER ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('FRESH DESIGN STUDIO')
    run.bold = True; run.font.size = Pt(11)
    run.font.color.rgb = GREEN

    add_para(doc, '', space_before=2, space_after=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Answer Engine Optimisation (AEO) Audit')
    run.bold = True; run.font.size = Pt(26)
    run.font.color.rgb = DARK_GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('bellanico.com')
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = GREEN

    add_para(doc, '', space_before=4, space_after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Prepared by: Fresh Design Studio   |   Date: {datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(10); run.font.color.rgb = MID_GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Confidential — Prepared exclusively for Bella Nico')
    run.font.size = Pt(9); run.font.color.rgb = MID_GREY; run.italic = True

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────
    add_heading(doc, '1. Executive Summary', level=1)
    add_divider(doc)
    add_para(doc,
        'Fresh Design Studio conducted a comprehensive Answer Engine Optimisation (AEO) audit of '
        'bellanico.com to evaluate how well the website is positioned to be discovered, understood, '
        'and cited by AI-powered answer engines including ChatGPT, Google AI Overviews, Perplexity, '
        'Claude, and Microsoft Copilot.',
        size=10, color=DARK_GREY, space_after=6)
    add_para(doc,
        'Bella Nico is a Canadian frozen food company established in 1995, celebrating 31 years in '
        'business in 2026. The company distributes premium frozen vegetables under the Veggie Values '
        'brand and frozen proteins including Roast Beef products. The website currently serves as a '
        'product catalogue and brand presence.',
        size=10, color=DARK_GREY, space_after=6)
    add_para(doc,
        'The audit identified several significant gaps that are preventing AI engines from accurately '
        'indexing, understanding, and recommending Bella Nico to consumers and buyers. The most '
        'critical issues — placeholder content on live pages, missing schema markup, and an '
        'unconfigured SEO plugin — require immediate attention. Addressing these issues will '
        'substantially improve how Bella Nico appears in AI-generated answers and search results.',
        size=10, color=DARK_GREY, space_after=6)

    # Score summary box
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    scores = [
        ('Schema\nMarkup', '0 / 10', RED),
        ('Meta &\nTitles',  '3 / 10', AMBER),
        ('Content\nQuality', '2 / 10', RED),
        ('Technical\nSEO',   '5 / 10', AMBER),
        ('Overall\nAEO Score', '3 / 10', RED),
    ]
    for i, (label, score, color) in enumerate(scores):
        cell = tbl.rows[0].cells[i]
        cell.width = Cm(3.0)
        set_cell_bg(cell, LIGHT_GREEN if i < 4 else LIGHT_GREY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(label + '\n')
        r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_GREY; r1.bold = True
        r2 = p.add_run(score)
        r2.font.size = Pt(18); r2.bold = True; r2.font.color.rgb = color

    doc.add_paragraph()

    # ── 2. WHAT IS AEO ────────────────────────────────────────────────────────
    add_heading(doc, '2. What Is Answer Engine Optimisation (AEO)?', level=1)
    add_divider(doc)
    add_para(doc,
        'Answer Engine Optimisation (AEO) is the practice of structuring your website\'s content, '
        'code, and authority signals so that AI-powered answer engines — including ChatGPT, Google '
        'AI Overviews, Perplexity AI, Microsoft Copilot, and Claude — can accurately understand, '
        'trust, and cite your business in their responses.',
        size=10, color=DARK_GREY, space_after=6)
    add_para(doc,
        'Unlike traditional SEO, which focuses on ranking in a list of blue links, AEO focuses on '
        'becoming the source that AI systems reference when answering questions such as:',
        size=10, color=DARK_GREY, space_after=4)
    for q in [
        '"Where can I buy frozen mini corn cobs in bulk?"',
        '"What brands offer club-pack frozen vegetables for distributors?"',
        '"Who makes Veggie Values frozen vegetables?"',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(q)
        run.italic = True; run.font.size = Pt(10); run.font.color.rgb = BLUE

    add_para(doc,
        'For Bella Nico, a company with a strong 31-year product history, AEO represents a significant '
        'untapped opportunity. The brand has authentic expertise and heritage that AI engines should '
        'be citing — but structural and content gaps are currently preventing that from happening.',
        size=10, color=DARK_GREY, space_before=6, space_after=6)

    # ── 3. SITE OVERVIEW ──────────────────────────────────────────────────────
    add_heading(doc, '3. Site Overview', level=1)
    add_divider(doc)

    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview = [
        ('Website',           'https://bellanico.com'),
        ('Company',           'Bella Nico Inc. — Veggie Values Brand'),
        ('Industry',          'Frozen Food Distribution / Wholesale'),
        ('WordPress Version', '6.9.4 (Latest — No Update Required)'),
        ('Active Theme',      'agroly-child 1.0 (Custom Child Theme)'),
        ('Active Plugins',    '28 Active | 3 Inactive | 16 Awaiting Updates'),
        ('Published Products','16 Products (Frozen Vegetables + Roast Beef Lines)'),
        ('Published Posts',   '0 — No Blog Content Published'),
    ]
    for i, (k, v) in enumerate(overview):
        row = tbl.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Cm(4.5); c1.width = Cm(11.5)
        set_cell_bg(c0, LIGHT_GREEN)
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]; run = p0.add_run(k)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = DARK_GREEN
        p1 = c1.paragraphs[0]; run = p1.add_run(v)
        run.font.size = Pt(9); run.font.color.rgb = DARK_GREY

    doc.add_paragraph()

    # ── 4. AEO AUDIT FINDINGS ────────────────────────────────────────────────
    add_heading(doc, '4. AEO Audit Findings', level=1)
    add_divider(doc)
    add_para(doc,
        'Each finding below is labelled to indicate where the fix needs to happen:',
        size=10, color=DARK_GREY, space_after=4)

    # Legend
    tbl_leg = doc.add_table(rows=1, cols=4)
    tbl_leg.alignment = WD_TABLE_ALIGNMENT.LEFT
    legends = [
        ('FRONT END', BLUE,     'Visible on the website — content, text, images'),
        ('BACK END',  DARK_GREY,'Behind the scenes — code, settings, plugins'),
    ]
    col = 0
    for label, color, desc in legends:
        c1 = tbl_leg.rows[0].cells[col];     c1.width = Cm(2.4)
        c2 = tbl_leg.rows[0].cells[col + 1]; c2.width = Cm(6.0)
        make_label_cell(c1, label, color)
        set_cell_bg(c2, LIGHT_GREY)
        p = c2.paragraphs[0]
        run = p.add_run(desc); run.font.size = Pt(9); run.font.color.rgb = MID_GREY
        col += 2
    doc.add_paragraph()

    # ── 4.1 CONTENT ──────────────────────────────────────────────────────────
    add_heading(doc, '4.1  Content Quality', level=2, color=DARK_GREEN, space_before=10)

    add_para(doc,
        'Content quality is the single most important factor in AEO. AI engines decide whether to '
        'trust, cite, and recommend a website based on the substance, accuracy, and specificity of '
        'its written content. The following content issues were identified on bellanico.com.',
        size=10, color=DARK_GREY, space_after=6)

    content_findings = [
        (
            'FAQ Page\n(Critical)',
            'The FAQ page at bellanico.com/faq contains Lorem Ipsum placeholder '
            'text — the same "dummy" content that came with the website theme. '
            'Questions include "What are gas solutions?" with fabricated answers '
            'that have no relation to Bella Nico\'s business.',
            'AI engines that crawl this page will either ignore it entirely or '
            'associate Bella Nico with irrelevant content. This actively harms '
            'brand authority with every AI crawler visit.',
            'CRITICAL',
            'FRONT END',
        ),
        (
            'Demo Contact\nDetails on FAQ',
            'The FAQ page displays the demo theme email addresses "info@agroly.com" '
            'and "suport@agroly.com" — these belong to the Agroly theme developer, '
            'not to Bella Nico. Visitors and AI crawlers see incorrect contact info.',
            'AI engines may associate Bella Nico with incorrect contact information, '
            'and real customers may attempt to contact the wrong email address.',
            'CRITICAL',
            'FRONT END',
        ),
        (
            'No Blog\nContent',
            'Zero blog posts have been published on the site. There is no '
            'article or informational content about frozen vegetables, food '
            'distribution, industry trends, or Bella Nico\'s expertise.',
            'AI engines cite pages that directly answer user questions. Without '
            'articles, Bella Nico cannot appear as a source for any industry '
            'query in ChatGPT, Perplexity, or Google AI Overviews.',
            'IMPORTANT',
            'FRONT END',
        ),
        (
            'Product\nDescriptions',
            '16 products are published (Veggie Values vegetables and Roast Beef '
            'lines) but most product pages lack rich descriptions, nutritional '
            'context, use cases, or distinguishing features.',
            'AI engines need written context to understand and recommend specific '
            'products. Sparse product pages cannot compete in AI-powered search.',
            'IMPORTANT',
            'FRONT END',
        ),
        (
            'Company\nAbout Page',
            'The "Bella Nico" page contains a strong brand history paragraph '
            '(established 1995, 31 years, first to introduce Mini Cob Corn) but '
            'lacks clear E-E-A-T signals: no leadership names, no certifications, '
            'no press coverage or industry credentials.',
            'AI engines weigh Experience, Expertise, Authoritativeness, and '
            'Trustworthiness (E-E-A-T). Strengthening this page raises the '
            'site\'s trust score with AI crawlers.',
            'RECOMMENDED',
            'FRONT END',
        ),
    ]
    add_finding_table(doc, content_findings)

    # ── 4.2 SCHEMA MARKUP ────────────────────────────────────────────────────
    add_heading(doc, '4.2  Schema Markup (Structured Data)', level=2, color=DARK_GREEN, space_before=10)
    add_para(doc,
        'Schema markup is code added to web pages that tells AI engines exactly what a piece of '
        'content represents — whether it\'s a product, a frequently asked question, a business, '
        'or an article. Without schema, AI engines must guess at the meaning of your content. '
        'With schema, you are telling them directly.',
        size=10, color=DARK_GREY, space_after=6)

    schema_findings = [
        (
            'No Schema\nConfigured',
            'AIOSEO (the SEO plugin installed on the site) has schema markup '
            'capability, but it has not been configured at all. The schema '
            'module is completely inactive.',
            'Every page on the site is missing structured data. AI engines '
            'receive no machine-readable context about who Bella Nico is, '
            'what products they sell, or what the site content represents.',
            'CRITICAL',
            'BACK END',
        ),
        (
            'No FAQ\nSchema',
            'The FAQ page content, even if rewritten with real questions and '
            'answers, will not benefit from AI indexing without FAQ schema '
            '(FAQPage + Question + Answer) markup.',
            'AI engines use FAQ schema to pull direct answers into their '
            'responses. Missing this schema means Bella Nico\'s FAQs will '
            'never appear as AI-cited answers.',
            'CRITICAL',
            'BACK END',
        ),
        (
            'No Product\nSchema',
            '16 product pages have no Product schema. There is no structured '
            'data for product name, description, category, or availability.',
            'Retailers, distributors, and consumers searching for specific '
            'frozen vegetable products will not find Bella Nico through AI '
            'product queries.',
            'CRITICAL',
            'BACK END',
        ),
        (
            'No Organisation\nSchema',
            'No Organisation schema exists to tell AI engines Bella Nico\'s '
            'official name, founding year (1995), industry, logo, and '
            'contact details.',
            'AI engines cannot confidently confirm basic facts about the '
            'company, reducing the likelihood of Bella Nico appearing in '
            '"who is" or "what is" brand queries.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            'No Breadcrumb\nSchema',
            'The site has a category/product hierarchy but no BreadcrumbList '
            'schema to communicate page relationships to AI crawlers.',
            'Without breadcrumb schema, AI engines struggle to understand '
            'the site\'s content architecture and how product pages relate '
            'to categories.',
            'RECOMMENDED',
            'BACK END',
        ),
    ]
    add_finding_table(doc, schema_findings)

    # ── 4.3 META & TITLES ────────────────────────────────────────────────────
    add_heading(doc, '4.3  Meta Titles & Descriptions', level=2, color=DARK_GREEN, space_before=10)
    add_para(doc,
        'Meta titles and descriptions are the first pieces of text AI engines and search engines '
        'read on every page. They act as a summary signal — if they are missing or generic, '
        'AI engines receive no clear context about page intent.',
        size=10, color=DARK_GREY, space_after=6)

    meta_findings = [
        (
            'Homepage\nMeta Missing',
            'The homepage has no custom meta title or description set in '
            'AIOSEO. The page is currently relying on the default WordPress '
            'site name only.',
            'The homepage is the most crawled page on the site. Without a '
            'descriptive meta title and description, AI engines lack the '
            'primary context signal about what Bella Nico does.',
            'CRITICAL',
            'BACK END',
        ),
        (
            'Most Pages\nNo Custom Meta',
            'With the exception of the Veggie Values page (which has a '
            'well-written title and description), all other pages — including '
            'Home, Contact, Careers, and individual product pages — are '
            'missing custom meta content.',
            'Each page without meta data is a missed opportunity for AI '
            'engines to understand and index that page\'s specific purpose.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            'Site Tagline\nRedundant Format',
            'The site\'s tagline is set to: "Bella Nico | Veggie Values | '
            'We took delicious and froze it in time". Using pipe separators '
            'in a tagline is a title format, not a meaningful description.',
            'When AI engines read this as a site description, it provides '
            'no informational value about the company\'s products, audience, '
            'or differentiators.',
            'RECOMMENDED',
            'BACK END',
        ),
    ]
    add_finding_table(doc, meta_findings)

    # ── 4.4 TECHNICAL ────────────────────────────────────────────────────────
    add_heading(doc, '4.4  Technical SEO & Crawlability', level=2, color=DARK_GREEN, space_before=10)
    add_para(doc,
        'Technical SEO covers the infrastructure that allows AI engines and search crawlers '
        'to discover, access, and index the site\'s content. Issues at this layer affect every '
        'page on the site simultaneously.',
        size=10, color=DARK_GREY, space_after=6)

    tech_findings = [
        (
            'No\nRobots.txt',
            'There is no robots.txt file at bellanico.com/robots.txt. '
            'This file provides instructions to all web crawlers — including '
            'AI training crawlers such as GPTBot, ClaudeBot, and PerplexityBot '
            '— about which pages they may or may not access.',
            'Without a robots.txt, AI training crawlers and search engine '
            'bots operate without guidance. The site cannot opt into or out '
            'of AI training datasets, and crawl budgets are unmanaged.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            'No XML\nSitemap',
            'AIOSEO has a sitemap feature but it has not been enabled. '
            'No XML sitemap exists at bellanico.com/sitemap.xml or any '
            'equivalent URL.',
            'Without a sitemap, AI engines and search crawlers may not '
            'discover all product pages, especially pages with limited '
            'internal links. Product pages risk being un-indexed.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            '436 Images\nMissing Alt Text',
            '436 uploaded images on the site have no alt text attribute. '
            'This includes product images, banner images, and content images '
            'across all pages.',
            'Alt text is how AI engines understand images. Missing alt text '
            'means hundreds of visual assets contribute zero context to the '
            'site\'s AI visibility. It also creates an accessibility '
            'compliance issue.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            '16 Plugin\nUpdates Pending',
            'Key plugins are significantly out of date, including: Elementor '
            '(3.25.6 → 4.0.3), Elementor Pro (3.25.5 → 4.0.2), ACF Pro '
            '(5.9.3 → 6.8.0.1), and AIOSEO (4.9.4.1 → 4.9.6.2). '
            'AIOSEO in particular may have AEO-specific improvements in '
            'newer versions.',
            'Outdated plugins can introduce security vulnerabilities, '
            'performance issues, and compatibility problems that indirectly '
            'harm crawlability and page speed scores.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            'AIOSEO Plugin\nNot Configured',
            'All In One SEO (AIOSEO) is installed but zero settings have '
            'been configured — including schema type, sitemap, social '
            'meta, and local business information.',
            'AIOSEO is one of the most powerful AEO tools available for '
            'WordPress. Having it inactive wastes its full capability. '
            'Every AEO gap in this report can be addressed through AIOSEO.',
            'CRITICAL',
            'BACK END',
        ),
        (
            'SSL\nConfigured',
            'The site is correctly served over HTTPS (SSL certificate active). '
            'This is a baseline requirement for any AI engine trust signal.',
            'Positive signal — no action required.',
            'GOOD',
            'BACK END',
        ),
        (
            'Caching\nActive',
            'SiteGround\'s SG CachePress caching plugin is active and '
            'running. Object caching is also enabled on the server.',
            'Positive signal — good page load performance foundation is '
            'in place. No action required for caching configuration.',
            'GOOD',
            'BACK END',
        ),
    ]
    add_finding_table(doc, tech_findings)

    # ── 4.5 AI VISIBILITY ─────────────────────────────────────────────────────
    add_heading(doc, '4.5  AI Visibility & Entity Recognition', level=2, color=DARK_GREEN, space_before=10)
    add_para(doc,
        'For a brand to be cited by AI engines, it must first be recognised as a distinct entity '
        '— a real, trustworthy business with verifiable attributes. The following factors affect '
        'how well AI engines can identify and represent Bella Nico as a brand.',
        size=10, color=DARK_GREY, space_after=6)

    entity_findings = [
        (
            'Brand Entity\nNot Established',
            'There is no page on the site that clearly presents Bella Nico '
            'as a formal entity: no structured company name in schema, '
            'no official social media profile links, no Wikipedia or '
            'knowledge panel presence.',
            'AI engines rely on named entity recognition to build knowledge '
            'graphs. Without entity signals, "Bella Nico" may not be '
            'consistently recognised as a distinct company versus a phrase.',
            'IMPORTANT',
            'BACK END',
        ),
        (
            'No Social\nProfiles Linked',
            'No social media profiles (LinkedIn, Facebook, Instagram) are '
            'linked from the website or declared in AIOSEO social settings.',
            'Social profile links in schema markup are a strong brand '
            'entity signal. AI engines cross-reference social presence '
            'when building brand knowledge.',
            'RECOMMENDED',
            'BACK END',
        ),
        (
            'Heritage &\nBrand Story',
            'The homepage contains valuable, authentic brand history: '
            'established 1995, first to introduce clear-pack Mini Cob Corn, '
            'celebrating 31 years in 2026. This is strong E-E-A-T content.',
            'This content is an asset — however it is currently mixed into '
            'Elementor layout blocks without semantic HTML markup, making '
            'it harder for AI engines to extract and cite cleanly.',
            'RECOMMENDED',
            'FRONT END',
        ),
        (
            'No Reviews\nor Testimonials',
            'No customer reviews, distributor testimonials, or industry '
            'certifications are presented on the site.',
            'Social proof and third-party endorsements are trust signals '
            'that AI engines factor into brand authority scores.',
            'RECOMMENDED',
            'FRONT END',
        ),
    ]
    add_finding_table(doc, entity_findings)

    doc.add_page_break()

    # ── 5. RECOMMENDATIONS ────────────────────────────────────────────────────
    add_heading(doc, '5. Recommendations', level=1)
    add_divider(doc)
    add_para(doc,
        'The following recommendations address every finding identified in this audit. '
        'Each item is labelled to show whether the work is visible on the website (Front End) '
        'or handled behind the scenes by a developer or website administrator (Back End).',
        size=10, color=DARK_GREY, space_after=8)

    recs = [
        (1,  'Replace FAQ Placeholder Content',
             'Rewrite the entire FAQ page with real Bella Nico questions and answers covering '
             'products, ordering, distribution, storage instructions, and company policies. '
             'Remove all Lorem Ipsum text. Update contact details to real Bella Nico information.',
             'FRONT END'),
        (2,  'Fix Demo Contact Information',
             'Remove "info@agroly.com" and "suport@agroly.com" from the FAQ page and any other '
             'pages where demo theme content remains. Replace with official Bella Nico contact '
             'details (phone, email, and address).',
             'FRONT END'),
        (3,  'Configure AIOSEO — Schema, Sitemap & Meta',
             'Set up AIOSEO with: (a) Organization schema for the company, (b) Product schema '
             'template for all product pages, (c) FAQPage schema for the FAQ page, (d) XML '
             'sitemap generation enabled and submitted to Google Search Console, (e) Social '
             'Open Graph settings, (f) custom meta titles and descriptions for all pages.',
             'BACK END'),
        (4,  'Create Homepage Meta Title & Description',
             'Write a homepage meta title (55–60 characters) and meta description (140–160 '
             'characters) that clearly describe Bella Nico, its products, and its unique value: '
             'e.g., "Bella Nico | Premium Frozen Vegetables & Club Pack Produce Since 1995".',
             'BACK END'),
        (5,  'Add Alt Text to All 436 Images',
             'Systematically add descriptive alt text to all 436 images missing this attribute. '
             'Use the AIOSEO image SEO feature or the WordPress Media Library. Focus first on '
             'all product images, then hero/banner images.',
             'BACK END'),
        (6,  'Create & Activate robots.txt',
             'Generate a robots.txt file via AIOSEO or manually. Include entries that allow '
             'major AI crawlers (GPTBot, ClaudeBot, PerplexityBot) to crawl the site, and '
             'disallow access to admin and internal WordPress directories.',
             'BACK END'),
        (7,  'Publish Product Descriptions',
             'Write unique, detailed descriptions for each of the 16 products. Include: '
             'product name, variety, pack size (e.g., 48oz), case count, key features, '
             'and typical use cases. This content directly enables AI product recommendations.',
             'FRONT END'),
        (8,  'Update Outdated Plugins',
             'Update all 16 plugins with available updates, prioritising: AIOSEO (SEO '
             'improvements), Elementor + Elementor Pro (major v4.0 release), and ACF Pro '
             '(5.9 → 6.8 is a significant update). Back up the site before updating.',
             'BACK END'),
        (9,  'Start a Blog / Resource Section',
             'Publish informational articles that answer real buyer and consumer questions: '
             '"Best frozen vegetables for meal prep", "How to read frozen vegetable packaging", '
             '"About our Veggie Values clear packaging". Target questions your audience asks '
             'AI engines.',
             'FRONT END'),
        (10, 'Link Social Media Profiles',
             'Add Bella Nico\'s LinkedIn, Facebook, and/or Instagram profile URLs to AIOSEO\'s '
             'Social settings. This helps AI engines establish Bella Nico as a verified brand '
             'entity across platforms.',
             'BACK END'),
        (11, 'Strengthen the About / Company Page',
             'Expand the Bella Nico brand story page to include: key milestones, named '
             'leadership (if applicable), distribution reach, any certifications or quality '
             'standards, and press or industry mentions.',
             'FRONT END'),
        (12, 'Update Site Tagline to a Real Description',
             'Replace the pipe-separated tagline "Bella Nico | Veggie Values | We took '
             'delicious and froze it in time" with a clear one-line description such as: '
             '"Bella Nico distributes premium frozen vegetables and proteins to retailers '
             'and distributors across North America."',
             'BACK END'),
    ]
    add_rec_table(doc, recs)

    # ── 6. ABOUT FDS ──────────────────────────────────────────────────────────
    add_heading(doc, '6. About Fresh Design Studio', level=1)
    add_divider(doc)
    add_para(doc,
        'Fresh Design Studio is a full-service digital agency specialising in AI-ready web '
        'design, Answer Engine Optimisation (AEO), and digital growth strategy for food, '
        'consumer goods, and distribution brands.',
        size=10, color=DARK_GREY, space_after=4)
    add_para(doc,
        'We help brands be found, understood, and cited by the AI engines that now drive '
        'consumer and B2B discovery decisions.',
        size=10, color=DARK_GREY, space_after=4)
    add_para(doc,
        'Contact: msakin@freshds.com',
        size=10, color=GREEN, bold=True, space_after=4)

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────
    add_para(doc, '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f'© {datetime.date.today().year} Fresh Design Studio  |  Confidential — Prepared for Bella Nico  |  bellanico.com'
    )
    run.font.size = Pt(8); run.font.color.rgb = MID_GREY; run.italic = True

    # ── SAVE ──────────────────────────────────────────────────────────────────
    out = Path(__file__).parent.parent / 'reports' / 'BellaNico_AEO_Audit_2026.docx'
    out.parent.mkdir(exist_ok=True)
    doc.save(str(out))
    print(f'Report saved: {out}')
    return str(out)


if __name__ == '__main__':
    generate()
